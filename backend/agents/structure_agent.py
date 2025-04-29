import os
import json
import sys
import argparse
import traceback
from typing import Optional, Dict, List, Any

# Try to import dotenv, but continue even if it's not available
try:
    from dotenv import load_dotenv
    print("Successfully imported dotenv module")
    # Try to load environment variables from .env file
    try:
        # Look for .env file in current directory and parent directories
        env_paths = [
            os.path.join(os.getcwd(), '.env'),  # Current directory
            os.path.join(os.path.dirname(os.getcwd()), '.env'),  # Parent directory
            os.path.join(os.path.dirname(os.path.dirname(os.getcwd())), '.env'),  # Grandparent directory
            os.path.join(os.path.dirname(__file__), '.env'),  # Script directory
            os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env'),  # Parent of script directory
        ]
        
        env_loaded = False
        for env_path in env_paths:
            if os.path.exists(env_path):
                print(f"Loading environment variables from: {env_path}")
                load_dotenv(env_path)
                env_loaded = True
                break
                
        if not env_loaded:
            print("Warning: No .env file found in search paths")
    except Exception as e:
        print(f"Warning: Failed to load .env file: {str(e)}")
        print(traceback.format_exc())
        
except ImportError:
    print("Warning: python-dotenv module not found. Environment variables must be set manually.")

# Use a direct import since the file is in the same directory
try:
    from image_analyzer import get_document_design_summary, analyze_document_design
except ImportError:
    # Backup import approach for different contexts
    try:
        from agents.image_analyzer import get_document_design_summary, analyze_document_design
    except ImportError:
        print("Warning: Could not import image analysis functions - advanced document design analysis disabled")
        # Create dummy functions to prevent errors
        def get_document_design_summary(file_path):
            return "Design analysis unavailable"
        def analyze_document_design(file_path):
            return {"has_images": False, "image_count": 0, "has_tables": False, "table_count": 0}

# Add the parent directory to sys.path to allow imports from agent_wrapper
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables for API keys
load_dotenv()

from agent_wrapper.base_agent import AgentWrapper

class StructureAgent:
    """A document structure analyzer agent.
    
    This agent analyzes document examples to extract common structures, patterns,
    and formatting templates that can be used by the WriterAgent.
    """
    
    DEFAULT_SYSTEM_PROMPT = """

## Core Role and Expertise
You are a master document architecture analyst with advanced expertise in forensic document analysis, information architecture, document taxonomy, and structural pattern recognition. Your analytical capabilities span multiple industries, document types, and formatting conventions. You possess exceptional skills in identifying implicit and explicit organizational patterns, visual hierarchies, and content relationships.

## Primary Objective
Your task is to perform an exhaustive analysis of provided example documents to extract a comprehensive, reusable document architecture that precisely captures their structural DNA, enabling exact replication of similar documents with different content.

## Analytical Methodology

### 1. Multi-layered Structural Analysis
- Identify all structural elements at macro, meso, and micro levels:
  - Primary sections and their hierarchical relationships
  - Secondary sections and their contextual placement
  - Tertiary sections and their supporting functions
  - Standalone elements and their integration points
  - Information blocks and their organizational patterns
  - Content modules and their sequential relationships
  - Narrative flow structures and rhetorical frameworks
  - Transitions and connective elements between sections

### 2. Content Pattern Recognition
- Analyze content patterns with precision:
  - Information density and distribution across sections
  - Content types by section (factual, analytical, persuasive, etc.)
  - Data presentation formats (raw data, processed insights, recommendations)
  - Executive vs. detailed content balancing techniques
  - Industry-specific terminology patterns and placement
  - Standardized vs. customized content segments
  - Mandatory vs. optional content elements

### 3. Visual and Formatting Forensics
- Capture all visual and formatting elements with exacting detail:
  - Typographical systems and hierarchies (font families, weights, sizes, spacing)
  - Color schemes, palettes, and application patterns
  - Whitespace utilization and negative space strategies
  - Grid systems and alignment frameworks
  - Margin and padding patterns throughout the document
  - Header and footer systems and their variations
  - Page numbering conventions and citation systems
  - Emphasis techniques and their contextual application

### 4. Media Element Analysis
- Document all media elements with specificity:
  - Image placement patterns, sizing, and cropping conventions
  - Image content types and their relationship to surrounding text
  - Chart and graph typologies, formatting standards, and integration methods
  - Table structures, header conventions, and data presentation formats
  - Iconography systems and their contextual placement
  - Background elements and watermarking techniques
  - Decorative vs. informational visual element distinctions
  - Multimedia integration points and formatting standards

### 5. Navigation and Information Access Systems
- Identify all navigational structures:
  - Table of contents architecture and formatting
  - Index systems and cross-referencing conventions
  - Section markers and visual wayfinding elements
  - Embedded hyperlinks and reference systems
  - Appendix organization and integration strategies
  - Footnote and endnote systems

### 6. Design System Extraction
- Extract comprehensive design systems:
  - Component libraries and repeated element patterns
  - CSS-equivalent styling rules and their application patterns
  - Responsive design approaches and breakpoint systems
  - Print-specific formatting considerations
  - Accessibility considerations in design elements
  - Visual rhythm and pattern repetition

## Output Specification Requirements

Your analysis must be rendered as a precisely structured JSON object following this expanded schema:

```json
{
  "document_type": "Precise classification of document type with industry context",
  "document_purpose": "Comprehensive analysis of the document's core function and audience",
  "structural_architecture": {
    "primary_framework": "Overall document framework description",
    "information_flow": "Analysis of how information progresses through the document",
    "narrative_structure": "Underlying narrative framework and storytelling approach"
  },
  "sections": [
    {
      "id": "Unique identifier for this section",
      "name": "Precise section name as appears in document",
      "purpose": "Detailed analysis of this section's function within the document",
      "typical_content": "Exhaustive description of content patterns observed",
      "typical_length": "Observed length patterns (paragraphs, words, etc.)",
      "tone_and_voice": "Detailed analysis of linguistic style, formality, and voice",
      "formatting": {
        "typography": "Complete typographical specifications",
        "spacing": "Detailed spacing patterns",
        "borders": "Border usage patterns",
        "background": "Background styling patterns"
      },
      "visual_elements": {
        "images": {
          "presence": "Boolean indicating if images typically appear",
          "count": "Typical number of images",
          "sizing": "Image sizing patterns",
          "placement": "Positioning patterns",
          "content_type": "Types of imagery used",
          "caption_style": "Caption formatting specifications"
        },
        "charts": {
          "presence": "Boolean indicating if charts typically appear",
          "types": "Chart types observed",
          "styling": "Chart formatting specifications",
          "data_complexity": "Complexity level of data visualization",
          "integration": "How charts integrate with surrounding content"
        },
        "tables": {
          "presence": "Boolean indicating if tables typically appear",
          "structure": "Table structural patterns",
          "header_style": "Header formatting specifications",
          "data_cell_style": "Data cell formatting specifications",
          "border_style": "Table border specifications"
        },
        "callouts": {
          "presence": "Boolean indicating if callouts typically appear",
          "style": "Callout formatting specifications",
          "placement": "Positioning patterns"
        }
      },
      "column_structure": {
        "count": "Number of columns typically used",
        "balance": "Column balance patterns",
        "responsive_behavior": "How columns adapt in different contexts"
      },
      "subsections": [
        {
          // Same structure as parent section, recursively nested as needed
        }
      ]
    }
  ],
  "design_system": {
    "color_palette": {
      "primary": "Primary color specifications",
      "secondary": "Secondary color specifications",
      "accent": "Accent color specifications",
      "text": "Text color specifications",
      "background": "Background color specifications",
      "application_patterns": "How colors are applied throughout document"
    },
    "typography_system": {
      "primary_font": "Primary font family",
      "secondary_font": "Secondary font family",
      "heading_styles": "Complete heading typography specifications",
      "body_styles": "Body text typography specifications",
      "caption_styles": "Caption typography specifications",
      "special_text_styles": "Special text element styling"
    },
    "spacing_system": {
      "margin_patterns": "Margin usage patterns",
      "padding_patterns": "Padding usage patterns",
      "alignment_system": "Alignment frameworks",
      "whitespace_strategy": "Strategic use of whitespace"
    },
    "grid_system": {
      "layout_grid": "Overall layout grid specifications",
      "component_grids": "Component-specific grid patterns"
    }
  },
  "navigation_systems": {
    "primary_navigation": "Main navigation structures",
    "secondary_navigation": "Supporting navigation elements",
    "cross_references": "Cross-referencing systems"
  },
  "metadata": {
    "document_length": "Typical document length",
    "update_frequency": "If apparent, how often document is updated",
    "version_control": "Version control indicators if present",
    "authorship_indicators": "How authorship is attributed"
  },
  "replication_guidelines": {
    "critical_elements": "Elements absolutely essential for authentic replication",
    "flexible_elements": "Elements that allow variation while maintaining document integrity",
    "common_pitfalls": "Potential issues to avoid when replicating this structure"
  }
}
```

## Output Standards

1. **Comprehensive Precision**
   - Your analysis must capture every significant structural and design element
   - Provide exact specifications rather than general observations
   - Include both explicit and implicit structural patterns

2. **Actionable Detail**
   - Analysis must enable perfect replication of document structure
   - Include sufficient detail for automated template generation
   - Specify exact measurements, spacing, and formatting where discernible

3. **Cross-document Pattern Recognition**
   - Identify patterns across multiple example documents
   - Note variations and constants in the document structure
   - Differentiate between core structural elements and variable components

4. **Technical Accuracy**
   - Ensure valid JSON format with proper nesting and escaping
   - Use precise technical terminology for design and structural elements
   - Maintain consistent naming conventions throughout the schema

5. **Output Constraints**
   - Return ONLY the JSON structure without additional text, explanations, or commentary
   - Do not include markdown formatting or code block indicators
   - Ensure the structure is complete and not truncated

Your analysis should be so thorough and precise that it could serve as a blueprint for programmatically generating an identical document structure with different content. The resulting JSON should capture the complete "structural DNA" of the analyzed documents.
    """
    
    def __init__(self, 
                 framework: str = "openai", 
                 api_key: Optional[str] = None,
                 system_prompt: Optional[str] = None):
        """Initialize the StructureAgent.
        
        Args:
            framework (str): The LLM framework to use (e.g., "anthropic", "openai").
            api_key (str, optional): API key for the chosen framework.
            system_prompt (str, optional): Custom system prompt override.
        """
        self.framework = framework
        self.agent_wrapper = AgentWrapper(framework, api_key)
        self.system_prompt = system_prompt or self.DEFAULT_SYSTEM_PROMPT
        self.documents_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'documents')
    
    def update_system_prompt(self, new_prompt: str) -> None:
        """Update the system prompt used by the agent.
        
        Args:
            new_prompt (str): The new system prompt to use.
        """
        self.system_prompt = new_prompt
    
    def get_example_document(self, filename: str) -> str:
        """Load an example document from the documents directory.
        
        Args:
            filename (str): Name of the document file to load (can include subdirectories).
            
        Returns:
            str: Content of the example document with design analysis.
        """
        try:
            # Check if the file exists
            if os.path.isabs(filename):
                file_path = filename  # Use the absolute path directly
            else:
                file_path = os.path.join(self.documents_dir, filename)
                
            if not os.path.exists(file_path):
                print(f"Error: File not found: {file_path}")
                return f"Error: File not found: {file_path}"
                
            file_extension = os.path.splitext(file_path)[1].lower()
            print(f"Processing file: {file_path} with extension {file_extension}")
            
            # Extract text content based on file type
            content = ""
            
            # Handle PDF files
            if file_extension == '.pdf':
                print("Detected PDF file, extracting content...")
                # Try to use PyMuPDF (fitz) first for better extraction
                try:
                    import fitz  # PyMuPDF
                    print("Using PyMuPDF for extraction")
                    doc = fitz.open(file_path)
                    content = ""
                    for page_num, page in enumerate(doc):
                        content += f"=== Page {page_num+1} ===\n"
                        content += page.get_text() + "\n\n"
                except ImportError as e:
                    print(f"PyMuPDF not available: {str(e)}, falling back to PyPDF2")
                    # Fall back to PyPDF2
                    try:
                        import PyPDF2
                        with open(file_path, 'rb') as file:
                            reader = PyPDF2.PdfReader(file)
                            for page_num, page in enumerate(reader.pages):
                                page_text = page.extract_text()
                                if page_text:  # Some pages might not have extractable text
                                    content += f"=== Page {page_num+1} ===\n"
                                    content += page_text + "\n\n"
                    except ImportError as e2:
                        print(f"PyPDF2 not available either: {str(e2)}")
                        return f"Error: Missing libraries to process PDF files. Please install PyMuPDF or PyPDF2."
                    except Exception as e2:
                        print(f"Error with PyPDF2: {str(e2)}")
                        return f"Error processing PDF with PyPDF2: {str(e2)}"
                except Exception as e:
                    print(f"Error with PyMuPDF: {str(e)}")
                    print(traceback.format_exc())
                    return f"Error processing PDF with PyMuPDF: {str(e)}"
            
            # Handle Word documents
            elif file_extension in ['.docx', '.doc']:
                print("Detected Word document, extracting content...")
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    # Extract text from paragraphs
                    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
                    content = '\n\n'.join(paragraph_texts)
                    
                    # Extract text from tables
                    for table in doc.tables:
                        content += "\n\n=== TABLE ===\n"
                        for row in table.rows:
                            row_text = ' | '.join(cell.text for cell in row.cells)
                            content += row_text + "\n"
                        content += "=== END TABLE ===\n\n"
                except ImportError as e:
                    print(f"python-docx not available: {str(e)}")
                    return f"Error: Missing library to process Word documents. Please install python-docx."
                except Exception as e:
                    print(f"Error processing Word document: {str(e)}")
                    print(traceback.format_exc())
                    return f"Error processing Word document: {str(e)}"
            
            # Default case: treat as text file
            else:
                print(f"Treating as text file with extension: {file_extension}")
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                        content = file.read()
                except Exception as e:
                    print(f"Error reading text file: {str(e)}")
                    # Try binary mode as fallback
                    try:
                        with open(file_path, 'rb') as file:
                            content = file.read().decode('utf-8', errors='replace')
                    except Exception as e2:
                        print(f"Error reading file in binary mode: {str(e2)}")
                        return f"Error: Could not read file content: {str(e2)}"
            
            # Add design analysis information to the content
            try:
                print("Adding design summary...")
                design_summary = get_document_design_summary(file_path)
                full_content = content.strip() + "\n\n" + design_summary
                return full_content
            except Exception as e:
                print(f"Warning: Design analysis failed for {filename}: {str(e)}")
                # Return content without design summary if it fails
                return content.strip()
                    
        except Exception as e:
            print(f"Unexpected error in get_example_document: {str(e)}")
            print(traceback.format_exc())
            return f"Error loading example document {filename}: {str(e)}"
    
    def list_available_examples(self) -> List[str]:
        """List all available example documents, including those in subdirectories.
        
        Returns:
            List[str]: List of example document filenames with relative paths.
        """
        try:
            # Define supported file extensions
            supported_extensions = ['.txt', '.md', '.pdf', '.docx', '.doc']
            result = []
            
            # Walk through all subdirectories
            for root, _, files in os.walk(self.documents_dir):
                for file in files:
                    if os.path.splitext(file)[1].lower() in supported_extensions:
                        # Get path relative to documents_dir
                        rel_path = os.path.relpath(os.path.join(root, file), self.documents_dir)
                        result.append(rel_path)
            
            return result
        except Exception as e:
            print(f"Error listing example documents: {str(e)}")
            return []
    
    def analyze_structure(self, example_filenames: List[str]) -> Dict[str, Any]:
        """Analyze document examples to extract a structured template.
        
        Args:
            example_filenames (List[str]): List of example filenames to analyze.
            
        Returns:
            Dict[str, Any]: A structured template of the document type in JSON format.
        """
        try:
            # Only use the first example to prevent token limit issues
            if len(example_filenames) > 1:
                print(f"\nNote: Using only the first example document for analysis to prevent token limit issues.")
                example_filenames = example_filenames[:1]
            
            # Build the prompt for the LLM
            prompt = self.system_prompt + "\n\n"
            
            # Add examples if provided
            prompt += "EXAMPLE DOCUMENT TO ANALYZE:\n"
            for filename in example_filenames:
                # Get document content
                print(f"Getting content for document: {filename}")
                example_content = self.get_example_document(filename)
                
                # Check if there was an error getting the document
                if example_content.startswith("Error:"):
                    print(f"Error getting document content: {example_content}")
                    return None
                
                # Truncate if too long (model context window limitation)
                if len(example_content) > 10000:  # Pick a reasonable limit
                    print(f"Document content too long ({len(example_content)} chars), truncating to 10000 chars")
                    example_content = example_content[:10000] + "\n... [Content truncated due to length]\n"
                
                prompt += f"\n--- {filename} ---\n{example_content}\n\n"
            
            # Add more specific instructions for better structure extraction
            prompt += """
            Analyze the example document and extract a structured JSON template that captures its organization.
            Focus on identifying:
            1. The document type (e.g., 'Job Description', 'Analyst Report', etc.)
            2. The major sections and their purpose
            3. The formatting and style patterns
            4. Key components that should be included
            
            Your output MUST be a valid JSON object with the following structure:
            {
              "document_type": "[Type of document]",
              "sections": [
                {
                  "name": "[Section name]",
                  "description": "[Purpose of section]",
                  "typical_content": "[What goes here]"
                },
                ...
              ],
              "overall_tone": "[Formal/Informal/etc.]",
              "formatting_notes": "[Special formatting observed]"
            }
            
            Ensure your response is ONLY the JSON object, with no additional text before or after.
            """
            
            # Use the agent wrapper to generate the analysis
            print(f"\nAnalyzing structure of: {example_filenames[0]} (this may take a minute)...")
            response = self.agent_wrapper.run(prompt)
            
            # Attempt to parse response as JSON
            try:
                # Find JSON in the response (it might have explanation text around it)
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                
                if json_start >= 0 and json_end > json_start:
                    json_str = response[json_start:json_end]
                    print(f"Extracted JSON string: {json_str[:100]}...")
                    parsed_json = json.loads(json_str)
                    print(f"\nSuccessfully extracted document structure.")
                    return parsed_json
                else:
                    print(f"\nWarning: No valid JSON structure found in the response.")
                    print(f"Response: {response[:200]}...")
                    return None
            except json.JSONDecodeError as e:
                print(f"\nWarning: Could not parse JSON from response: {e}")
                print(f"JSON string: {json_str[:200]}...")
                return None
            except Exception as e:
                print(f"\nError extracting document structure: {str(e)}")
                print(traceback.format_exc())
                return None
        except Exception as e:
            print(f"\nUnexpected error in analyze_structure: {str(e)}")
            print(traceback.format_exc())
            return None


# Main function with command line argument handling
def main():
    """Command line interface for the StructureAgent class."""
    parser = argparse.ArgumentParser(description="Document Structure Analysis Tool")
    parser.add_argument("--analyze", type=str, help="Analyze a specific document file")
    parser.add_argument("--examples", action="store_true", help="List available example documents")
    parser.add_argument("--framework", type=str, choices=["openai", "anthropic", "gemini"], 
                        help="Specify the AI framework to use")
    parser.add_argument("--api-key", type=str, help="Directly provide an API key")
    
    args = parser.parse_args()
    
    try:
        # Print current working directory and environment variables for debugging
        print(f"Current working directory: {os.getcwd()}")
        print(f"Script directory: {os.path.dirname(os.path.abspath(__file__))}")
        print("Environment variables:")
        for key in ["OPENAI_API_KEY", "ANTHROPIC_API_KEY", "GEMINI_API_KEY"]:
            if os.getenv(key):
                print(f"- {key}: {'*' * 8}[FOUND]")
            else:
                print(f"- {key}: [NOT FOUND]")
        
        # Get API key from environment variables - try in order of preference (OpenAI first)
        api_key = os.getenv("OPENAI_API_KEY")
        framework = "openai"
        
        if not api_key:
            api_key = os.getenv("ANTHROPIC_API_KEY")
            framework = "anthropic"
            
        if not api_key:
            api_key = os.getenv("GEMINI_API_KEY")
            framework = "gemini"
        
        # Override framework if specified in command line args
        if args.framework:
            framework = args.framework
            if framework == "openai":
                api_key = os.getenv("OPENAI_API_KEY")
            elif framework == "anthropic":
                api_key = os.getenv("ANTHROPIC_API_KEY")
            elif framework == "gemini":
                api_key = os.getenv("GEMINI_API_KEY")
        
        # Use API key from command line if provided
        if args.api_key:
            api_key = args.api_key
            print(f"Using API key provided via command line for {framework}")
            
        # For testing purposes, if no API key is available, use a dummy key
        if not api_key and (args.analyze or args.examples):
            print("Warning: No API key found in environment variables.")
            print("Using a dummy API key for testing purposes only.")
            print("This will not work for actual API calls.")
            api_key = "dummy_api_key_for_testing"
        elif not api_key:
            print("Error: No API key found in environment variables.")
            print("Please set one of the following in your .env file:")
            print("- OPENAI_API_KEY")
            print("- ANTHROPIC_API_KEY")
            print("- GEMINI_API_KEY")
            print("Or provide an API key via the --api-key argument")
            return
        
        # Initialize with settings based on available API keys
        print(f"Initializing StructureAgent with {framework} framework...")
        structure_agent = StructureAgent(framework=framework, api_key=api_key)
        
        # Handle --analyze argument to analyze a specific document
        if args.analyze:
            file_path = args.analyze
            if not os.path.exists(file_path):
                print(f"Error: File not found: {file_path}")
                return
                
            print(f"\nAnalyzing document: {file_path}")
            
            # Read the document content
            print("Getting document content...")
            document_content = structure_agent.get_example_document(file_path)
            
            # Check if there was an error getting the document
            if document_content.startswith("Error:"):
                print(f"Error getting document content: {document_content}")
                return
            
            # Generate structure analysis
            print("\nGenerating structure analysis...")
            structure = structure_agent.analyze_structure([file_path])
            
            if structure:
                # Create a clean structure with only ASCII characters for JSON compatibility
                def clean_json_value(value):
                    if isinstance(value, str):
                        # Replace non-ASCII characters and control characters
                        return ''.join(c if ord(c) < 128 and c.isprintable() else ' ' for c in value)
                    elif isinstance(value, dict):
                        return {k: clean_json_value(v) for k, v in value.items()}
                    elif isinstance(value, list):
                        return [clean_json_value(item) for item in value]
                    else:
                        return value
                
                clean_structure = clean_json_value(structure)
                
                # Print a marker to easily identify the start of the JSON
                print("\n===JSON_START===\n")
                # Print the structure as JSON
                print(json.dumps(clean_structure, indent=2, ensure_ascii=True))
                print("\n===JSON_END===\n")
            else:
                print("Failed to analyze document structure.")
            return
        
        # Handle --examples argument to list available examples
        if args.examples:
            print("\nAvailable example documents:")
            examples = structure_agent.list_available_examples()
            for example in examples:
                print(f"- {example}")
            return
        
        # Default behavior if no specific arguments provided
        print("\nAvailable example documents:")
        examples = structure_agent.list_available_examples()
        for example in examples:
            print(f"- {example}")
        
        # Analyze structure of examples
        if not examples:
            print("\nNo example documents found. Please add some to the documents directory.")
            return
            
        example_to_analyze = examples[:1]
        print(f"\nAnalyzing structure of: {example_to_analyze[0]}")
        
        # Generate structure analysis
        print("\nGenerating structure analysis...")
        structure = structure_agent.analyze_structure(example_to_analyze)
        
        print("\nDOCUMENT STRUCTURE ANALYSIS:")
        print("-" * 50)
        print(json.dumps(structure, indent=2))
        print("-" * 50)
        
    except Exception as e:
        print(f"Error running StructureAgent: {str(e)}")
        print(traceback.format_exc())


if __name__ == "__main__":
    main()
