"""
Image analysis helper module for StructureAgent.

This module provides functions to extract and analyze images from documents,
helping the StructureAgent better understand document design elements.
"""

import os
import io
from typing import List, Dict, Any
from PIL import Image

# For PDF processing
import fitz  # PyMuPDF

def extract_images_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Extract images from a PDF file with position information.
    
    Args:
        pdf_path (str): Path to the PDF file.
        
    Returns:
        List[Dict[str, Any]]: List of dictionaries containing image data and metadata.
    """
    images = []
    
    try:
        # Open the PDF
        doc = fitz.open(pdf_path)
        
        # Iterate through each page
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            
            # Process each image on the page
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Get position information (from the image rectangle)
                try:
                    # Find the image rectangle on the page
                    img_rect = None
                    for rect in page.get_image_bbox(xref):
                        img_rect = rect
                        break
                    
                    if img_rect:
                        position = {
                            "x": img_rect.x0,
                            "y": img_rect.y0,
                            "width": img_rect.width,
                            "height": img_rect.height,
                            "page_width": page.rect.width,
                            "page_height": page.rect.height,
                        }
                    else:
                        position = None
                except Exception:
                    position = None
                
                # Try to get image dimensions
                try:
                    img = Image.open(io.BytesIO(image_bytes))
                    width, height = img.size
                    format_name = img.format
                except Exception:
                    width, height, format_name = None, None, None
                
                # Store the image data
                image_data = {
                    "page_num": page_num + 1,  # 1-based page number
                    "image_index": img_index,
                    "position": position,
                    "width": width,
                    "height": height,
                    "format": format_name or image_ext,
                    "size_bytes": len(image_bytes),
                }
                
                images.append(image_data)
        
        return images
    
    except Exception as e:
        print(f"Error extracting images from PDF: {str(e)}")
        return []

def extract_images_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract images from a DOCX file with position information.
    
    Args:
        docx_path (str): Path to the DOCX file.
        
    Returns:
        List[Dict[str, Any]]: List of dictionaries containing image data and metadata.
    """
    images = []
    
    try:
        import docx
        from docx.document import Document
        from docx.parts.image import ImagePart
        
        # Open the DOCX file
        doc = docx.Document(docx_path)
        
        # Helper function to process individual runs/paragraphs
        def process_paragraph(paragraph, para_index):
            for run_index, run in enumerate(paragraph.runs):
                # Check if run contains an image
                if hasattr(run, "_element") and run._element.findall(".//a:blip", namespaces=run._element.nsmap):
                    # Extract image info - a bit complex in python-docx
                    embedded_images = []
                    
                    # Look for relationships that are images
                    for rel in paragraph.part.rels.values():
                        if isinstance(rel._target, ImagePart):
                            # We found an image
                            try:
                                # Get image binary data
                                image_bytes = rel._target.blob
                                
                                # Try to get image dimensions
                                try:
                                    img = Image.open(io.BytesIO(image_bytes))
                                    width, height = img.size
                                    format_name = img.format
                                except Exception:
                                    width, height, format_name = None, None, None
                                
                                # Store image data
                                image_data = {
                                    "paragraph_index": para_index,
                                    "run_index": run_index,
                                    "width": width,
                                    "height": height,
                                    "format": format_name,
                                    "size_bytes": len(image_bytes),
                                    # Position is approximate in DOCX as exact coordinates aren't available
                                    "position": None  
                                }
                                
                                embedded_images.append(image_data)
                            except Exception as e:
                                print(f"Error processing image in DOCX: {str(e)}")
                    
                    images.extend(embedded_images)
        
        # Process all paragraphs
        for para_index, paragraph in enumerate(doc.paragraphs):
            process_paragraph(paragraph, para_index)
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para_index, paragraph in enumerate(cell.paragraphs):
                        process_paragraph(paragraph, para_index)
        
        return images
    
    except Exception as e:
        print(f"Error extracting images from DOCX: {str(e)}")
        return []

def extract_tables_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Extract tables from a PDF file.
    
    Args:
        pdf_path (str): Path to the PDF file.
        
    Returns:
        List[Dict[str, Any]]: List of dictionaries containing table metadata.
    """
    tables = []
    
    try:
        # Open the PDF
        doc = fitz.open(pdf_path)
        
        # Iterate through each page
        for page_num, page in enumerate(doc):
            # Simple table detection using text blocks
            blocks = page.get_text("dict")["blocks"]
            
            # Group blocks that might be part of tables (looking for grid-like structures)
            potential_tables = []
            
            # Basic algorithm to detect tables:
            # 1. Find blocks that seem to be aligned in a grid pattern
            # 2. Group these blocks together as potential tables
            
            # This is a simplified approach - for production use, consider specialized table extraction libraries
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    lines = block.get("lines", [])
                    if len(lines) > 1:
                        # Check for consistent line spacing and similar widths (table-like characteristics)
                        line_heights = [line.get("bbox")[3] - line.get("bbox")[1] for line in lines]
                        line_widths = [line.get("bbox")[2] - line.get("bbox")[0] for line in lines]
                        
                        avg_height = sum(line_heights) / len(line_heights) if line_heights else 0
                        height_variation = max(line_heights) - min(line_heights) if line_heights else 0
                        
                        # If lines have consistent height and more than 2 lines, it might be a table
                        if len(lines) > 2 and height_variation < avg_height * 0.2:
                            cells = []
                            for line in lines:
                                for span in line.get("spans", []):
                                    cells.append({
                                        "text": span.get("text", ""),
                                        "bbox": span.get("bbox", [0, 0, 0, 0])
                                    })
                            
                            potential_tables.append({
                                "page_num": page_num + 1,
                                "bbox": block.get("bbox", [0, 0, 0, 0]),
                                "rows": len(lines),
                                "cells": cells
                            })
            
            tables.extend(potential_tables)
        
        return tables
    
    except Exception as e:
        print(f"Error extracting tables from PDF: {str(e)}")
        return []

def analyze_document_design(file_path: str) -> Dict[str, Any]:
    """
    Analyze a document's design elements including images, tables, and layout.
    
    Args:
        file_path (str): Path to the document file.
        
    Returns:
        Dict[str, Any]: Dictionary containing design analysis results.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    design_info = {
        "has_images": False,
        "image_count": 0,
        "images": [],
        "has_tables": False, 
        "table_count": 0,
        "tables": [],
        "layout_elements": []
    }
    
    # Process based on file type
    if file_extension == '.pdf':
        # Extract images
        images = extract_images_from_pdf(file_path)
        design_info["images"] = images
        design_info["image_count"] = len(images)
        design_info["has_images"] = len(images) > 0
        
        # Extract tables
        tables = extract_tables_from_pdf(file_path)
        design_info["tables"] = tables
        design_info["table_count"] = len(tables)
        design_info["has_tables"] = len(tables) > 0
        
        # Add layout information (simplified)
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                layout_data = {
                    "page_num": page_num + 1,
                    "width": page.rect.width,
                    "height": page.rect.height,
                    "text_blocks": len(page.get_text("dict")["blocks"])
                }
                design_info["layout_elements"].append(layout_data)
        except Exception as e:
            print(f"Error analyzing PDF layout: {str(e)}")
    
    elif file_extension in ['.docx', '.doc']:
        # Extract images
        images = extract_images_from_docx(file_path)
        design_info["images"] = images
        design_info["image_count"] = len(images)
        design_info["has_images"] = len(images) > 0
        
        # Get basic layout info
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Count tables
            table_count = len(doc.tables)
            design_info["table_count"] = table_count
            design_info["has_tables"] = table_count > 0
            
            # Simplified table info
            tables = []
            for i, table in enumerate(doc.tables):
                tables.append({
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0
                })
            design_info["tables"] = tables
            
            # Basic layout info
            design_info["layout_elements"] = [{
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
                "tables": table_count
            }]
        except Exception as e:
            print(f"Error analyzing DOCX layout: {str(e)}")
    
    return design_info

def get_document_design_summary(file_path: str) -> str:
    """
    Generate a text summary of document design elements for LLM analysis.
    
    Args:
        file_path (str): Path to the document file.
        
    Returns:
        str: Text summary of document design.
    """
    design_info = analyze_document_design(file_path)
    
    summary_lines = [
        "DOCUMENT DESIGN ANALYSIS:",
        "======================="
    ]
    
    # Image summary
    if design_info["has_images"]:
        summary_lines.append(f"IMAGES: {design_info['image_count']} images detected")
        for i, img in enumerate(design_info["images"][:5]):  # Limit to first 5 images
            page_info = f"Page {img['page_num']}" if 'page_num' in img else "Document"
            size_info = f"{img.get('width', 'unknown')}x{img.get('height', 'unknown')}" if img.get('width') else "Unknown size"
            summary_lines.append(f"  Image {i+1}: {page_info}, {size_info}, Format: {img.get('format', 'unknown')}")
        if design_info["image_count"] > 5:
            summary_lines.append(f"  ... and {design_info['image_count'] - 5} more images")
    else:
        summary_lines.append("IMAGES: No images detected")
    
    # Table summary
    if design_info["has_tables"]:
        summary_lines.append(f"TABLES: {design_info['table_count']} tables detected")
        for i, table in enumerate(design_info["tables"][:3]):  # Limit to first 3 tables
            if "rows" in table:
                summary_lines.append(f"  Table {i+1}: {table.get('rows', 'unknown')} rows, {table.get('columns', 'unknown')} columns")
            elif "page_num" in table:
                summary_lines.append(f"  Table {i+1}: On page {table.get('page_num')}, approximately {len(table.get('cells', []))} cells")
        if design_info["table_count"] > 3:
            summary_lines.append(f"  ... and {design_info['table_count'] - 3} more tables")
    else:
        summary_lines.append("TABLES: No tables detected")
    
    # Layout summary
    summary_lines.append("LAYOUT:")
    for layout in design_info["layout_elements"]:
        if "page_num" in layout:
            summary_lines.append(f"  Page {layout['page_num']}: {layout.get('width', 'unknown')}x{layout.get('height', 'unknown')}, {layout.get('text_blocks', 'unknown')} text blocks")
        else:
            summary_lines.append(f"  Document has {layout.get('paragraphs', 'unknown')} paragraphs, {layout.get('sections', 'unknown')} sections")
    
    return "\n".join(summary_lines)