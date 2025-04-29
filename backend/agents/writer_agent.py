import os
import json
import sys
from typing import Optional, Dict, List, Any
from dotenv import load_dotenv

# Add the parent directory to sys.path to allow imports from agent_wrapper
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables for API keys
load_dotenv()

from agent_wrapper.base_agent import AgentWrapper

class WriterAgent:
    """A document writer agent that leverages LLM capabilities.
    
    This agent uses the agent wrapper to interact with various LLM providers
    and generates documents based on example data and user inputs.
    """
    
    DEFAULT_SYSTEM_PROMPT = """
    You are an expert document writer with expertise in creating professional documents.
    Your task is to create new documents based on the examples provided and the user's specifications.
    
    Follow these guidelines:
    1. Analyze the example documents to understand their structure, style, and tone.
    2. Create a new document that follows a similar format but is tailored to the user's requirements.
    3. Ensure the content is clear, concise, and professionally written.
    4. Maintain the appropriate level of formality based on the document type.
    5. Include all necessary sections and details typically found in this type of document.
    6. Create the content as pure HTML code, with no Python variable references.
    7. ALWAYS write full HTML as one continuous literal string, properly escaped in Python.
    8. NEVER attempt to use variable substitution for any HTML or CSS elements.
    9. CSS properties like 'border', 'margin', 'padding', etc. should ONLY appear within string literals.
    10. Do not define variables for HTML/CSS properties like: border = "2px solid black".
    11. For HTML output, always ensure the format is a complete single string with no embedded variables.
    12. Use the example structure provided as a guide for content and organization.
    13. If writing HTML, ensure all attributes and CSS properties are properly quoted inside the string.
    
    CRITICAL: Your entire output must be returnable as a single string with no variable references. 
    Simply generate the complete HTML document as a direct literal string.
    
    Be creative while staying professional and relevant to the user's needs.
    """
    
    def __init__(self, 
                 framework: str = "openai", 
                 api_key: Optional[str] = None,
                 system_prompt: Optional[str] = None):
        """Initialize the WriterAgent.
        
        Args:
            framework (str): The LLM framework to use (e.g., "anthropic", "openai").
            api_key (str, optional): API key for the chosen framework.
            system_prompt (str, optional): Custom system prompt override.
        """
        self.framework = framework
        self.agent_wrapper = AgentWrapper(framework, api_key)
        self.system_prompt = system_prompt or self.DEFAULT_SYSTEM_PROMPT
        
        # Set directories for examples and knowledge base
        self.base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.documents_dir = os.path.join(self.base_dir, 'documents')
        self.knowledge_base_dir = os.path.join(self.base_dir, 'knowledge_base')
    
    def update_system_prompt(self, new_prompt: str) -> None:
        """Update the system prompt used by the agent.
        
        Args:
            new_prompt (str): The new system prompt to use.
        """
        self.system_prompt = new_prompt
    
    def get_example_document(self, filename: str) -> str:
        """Load an example document from the documents directory.
        
        Args:
            filename (str): Name of the document file to load (can include subdirectories).
            
        Returns:
            str: Content of the example document.
        """
        try:
            file_path = os.path.join(self.documents_dir, filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Handle PDF files
            if file_extension == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:  # Some pages might not have extractable text
                            text += page_text + "\n\n"
                    return text.strip()
            
            # Handle Word documents
            elif file_extension in ['.docx', '.doc']:
                import docx
                doc = docx.Document(file_path)
                return '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            
            # Default case: treat as text file
            else:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                    return file.read()
                    
        except ImportError as e:
            return f"Error: Missing library to process {filename}. {str(e)}"
        except Exception as e:
            return f"Error loading example document {filename}: {str(e)}"
    
    def list_available_examples(self) -> List[str]:
        """List all available example documents, including those in subdirectories.
        
        Returns:
            List[str]: List of example document filenames with relative paths.
        """
        try:
            # Define supported file extensions
            supported_extensions = ['.txt', '.md', '.pdf', '.docx', '.doc']
            result = []
            
            # Walk through all subdirectories
            for root, _, files in os.walk(self.documents_dir):
                for file in files:
                    if os.path.splitext(file)[1].lower() in supported_extensions:
                        # Get path relative to documents_dir
                        rel_path = os.path.relpath(os.path.join(root, file), self.documents_dir)
                        result.append(rel_path)
            
            return result
        except Exception as e:
            print(f"Error listing example documents: {str(e)}")
            return []
    
    def load_knowledge_base_content(self) -> Dict[str, str]:
        """Load content from all files in the knowledge_base directory.
        
        Returns:
            Dict[str, str]: Dictionary mapping filenames to their content.
        """
        knowledge_base = {}
        
        try:
            if os.path.exists(self.knowledge_base_dir) and os.path.isdir(self.knowledge_base_dir):
                # Define supported file extensions
                supported_extensions = ['.txt', '.md', '.pdf', '.docx', '.doc']
                
                # Walk through all subdirectories
                for root, _, files in os.walk(self.knowledge_base_dir):
                    for file in files:
                        if os.path.splitext(file)[1].lower() in supported_extensions:
                            # Get path relative to knowledge_base_dir
                            rel_path = os.path.relpath(os.path.join(root, file), self.knowledge_base_dir)
                            full_path = os.path.join(root, file)
                            
                            # Get file extension
                            file_extension = os.path.splitext(file)[1].lower()
                            content = ""
                            
                            # Process file based on extension
                            try:
                                if file_extension == '.pdf':
                                    import PyPDF2
                                    with open(full_path, 'rb') as f:
                                        reader = PyPDF2.PdfReader(f)
                                        text = ""
                                        for page in reader.pages:
                                            page_text = page.extract_text()
                                            if page_text:
                                                text += page_text + "\n\n"
                                        content = text.strip()
                                elif file_extension in ['.docx', '.doc']:
                                    import docx
                                    doc = docx.Document(full_path)
                                    content = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
                                else:
                                    with open(full_path, 'r', encoding='utf-8', errors='replace') as f:
                                        content = f.read()
                                        
                                knowledge_base[rel_path] = content
                            except Exception as e:
                                print(f"Error loading knowledge base file {rel_path}: {str(e)}")
        except Exception as e:
            print(f"Error loading knowledge base: {str(e)}")
            
        return knowledge_base
    
    def create_document(self, user_input: str, example_filenames: Optional[List[str]] = None) -> str:
        """Create a new document based on user input and examples.
        
        Args:
            user_input (str): User's requirements for the document.
            example_filenames (List[str], optional): List of example filenames to reference.
            
        Returns:
            str: The generated document.
        """
        # Build the prompt for the LLM
        prompt = self.system_prompt + "\n\n"
        
        # Add examples if provided
        if example_filenames:
            prompt += "EXAMPLE DOCUMENTS:\n"
            for filename in example_filenames:
                example_content = self.get_example_document(filename)
                prompt += f"\n--- {filename} ---\n{example_content}\n\n"
        
        # Add user's requirements
        prompt += "USER REQUIREMENTS:\n" + user_input + "\n\n"
        prompt += "Please create a new document based on the examples and user requirements:"
        
        # Use the agent wrapper to generate the document
        return self.agent_wrapper.run(prompt)
    
    def create_document_with_structure(self, user_input: str, structure: Dict[str, Any], example_filenames: Optional[List[str]] = None) -> str:
        """Create a new document based on user input, a predefined structure, and knowledge base data.
        
        Args:
            user_input (str): User's requirements for the document.
            structure (Dict[str, Any]): Document structure template to follow.
            example_filenames (List[str], optional): List of example filenames to reference.
            
        Returns:
            str: The generated document.
        """
        # Build the prompt for the LLM
        prompt = self.system_prompt + "\n\n"
        
        # Add document structure
        prompt += "DOCUMENT STRUCTURE TEMPLATE:\n"
        prompt += json.dumps(structure, indent=2) + "\n\n"
        prompt += """
        Follow this document structure template carefully when creating the new document.
        The structure template is a JSON object that defines the sections and subsections of the document.
        Output the document in markdown format.
        Use the following markdown formatting: 
        * Use backticks for inline code
        * Use double backticks for code blocks
        * Use asterisks for italics
        * Use double asterisks for bold
        * Use tildes for strikethrough
        * Use colons for horizontal rules
        * Use angle brackets for URLs
        * Use backticks for inline code
        * Use double backticks for code blocks
        * Use asterisks for italics
        * Use double asterisks for bold
        * Use tildes for strikethrough
        * Use colons for horizontal rules
        * Use angle brackets for URLs
        \n\n
        """
        
        # Load and incorporate knowledge base content
        knowledge_base = self.load_knowledge_base_content()
        if knowledge_base:
            prompt += "KNOWLEDGE BASE CONTENT:\n"
            for filename, content in knowledge_base.items():
                # Truncate very long files to avoid hitting token limits
                if len(content) > 5000:
                    displayed_content = content[:5000] + "\n[Content truncated due to length...]\n"
                else:
                    displayed_content = content
                    
                prompt += f"\n--- {filename} ---\n{displayed_content}\n"
            
            prompt += "\n\nUse the knowledge base content above as the source of factual information for creating the document.\n\n"
        
        # Add examples if provided (for reference, but knowledge base and structure take precedence)
        if example_filenames:
            prompt += "REFERENCE DOCUMENTS:\n"
            for filename in example_filenames:
                example_content = self.get_example_document(filename)
                prompt += f"\n--- {filename} ---\n{example_content}\n\n"
            prompt += "Use these documents as additional references.\n\n"
        
        # Add user's requirements
        prompt += "USER REQUIREMENTS:\n" + user_input + "\n\n"
        prompt += "Please create a new document based on the structure template, user requirements, and Knowledge Base. Focus on using factual information from the Knowledge Base while following the document structure and addressing user requirements."
        
        # Use the agent wrapper to generate the document
        return self.agent_wrapper.run(prompt)


def main():
    """Example usage of the WriterAgent class."""
    try:
        # Get API key from environment variables - try in order of preference (OpenAI first)
        api_key = os.getenv("OPENAI_API_KEY")
        framework = "openai"
        
        if not api_key:
            api_key = os.getenv("ANTHROPIC_API_KEY") 
            framework = "anthropic"
            
        if not api_key:
            api_key = os.getenv("GEMINI_API_KEY")
            framework = "gemini"
            
        if not api_key:
            print("Error: No API key found in environment variables.")
            print("Please set one of the following in your .env file:")
            print("- OPENAI_API_KEY")
            print("- ANTHROPIC_API_KEY")
            print("- GEMINI_API_KEY")
            return
        
        # Initialize with settings based on available API keys
        print(f"Initializing WriterAgent with {framework} framework...")
        writer = WriterAgent(framework=framework, api_key=api_key)
        
        # List available example documents
        print("\nAvailable example documents:")
        examples = writer.list_available_examples()
        for example in examples:
            print(f"- {example}")
        
        # Example user input
        user_input = """
        Create a job description for a Senior Data Scientist role at XYZ Analytics.
        The position requires 5+ years of experience, expertise in machine learning and Python,
        and offers a salary range of $140,000-$180,000.
        """
        
        # Generate document using the first example (if available)
        if not examples:
            print("\nNo example documents found. Please add some to the documents directory.")
            return
            
        example_to_use = examples[:1]
        print(f"\nUsing example document: {example_to_use[0]}")
        
        # Generate document
        print("\nGenerating document...")
        generated_document = writer.create_document(user_input, example_to_use)
        
        print("\nGENERATED DOCUMENT:")
        print("-" * 50)
        print(generated_document)
        print("-" * 50)
        
        # Example of updating the system prompt
        print("\nUpdating system prompt...")
        custom_prompt = """
        You are a creative technical writer specializing in job descriptions.
        Create engaging and detailed job descriptions that accurately represent the role
        while attracting top talent. Focus on company culture and growth opportunities.
        """
        writer.update_system_prompt(custom_prompt)
        
        # Generate another document with the updated prompt
        print("\nGenerating document with updated prompt...")
        new_document = writer.create_document(user_input, example_to_use)
        
        print("\nGENERATED DOCUMENT WITH UPDATED PROMPT:")
        print("-" * 50)
        print(new_document)
        print("-" * 50)
        
    except Exception as e:
        print(f"Error running WriterAgent: {str(e)}")


if __name__ == "__main__":
    main()